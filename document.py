import os

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PyPDF2 import PdfReader
from docx2pdf import convert
from rich.progress import Progress

from bilingual import BilingualCreater


class DocumentCreater(BilingualCreater):

    def __init__(self):
        self.pars_per_page = []


    def create_bilingual_doc(self):
        """  """
        self.determine_paragraph_counts()
        
        tr_par = self.tr_par.copy()
        en_par = self.en_par.copy()

        # Create a new A5 document
        doc = self.create_doc()
        
        with Progress(transient=True) as progress:
            task = progress.add_task("[green]Creating bilingual document...", total=len(self.pars_per_page))

            for par_num in self.pars_per_page:
                
                # Turkish page
                doc.add_page_break()
                for x in range(par_num):
                    if tr_par:
                        par = doc.add_paragraph(tr_par.pop(0))
                        par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                # English page
                doc.add_page_break()
                for x in range(par_num):
                    if en_par:
                        par = doc.add_paragraph(en_par.pop(0))
                        par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                # Update progress
                progress.update(task, advance=1, description="[green]Creating bilingual document...")
                 
            doc.save("bilingual_book.docx")
        
        self.remove_temps()


    def determine_paragraph_counts(self):
        """  """
        tr_par = self.tr_par.copy()
        en_par = self.en_par.copy()

        progress_len = len(en_par)
        
        with Progress(transient=True) as progress:
            task = progress.add_task("[green]Calculating paragraph counts per page...", total=progress_len)

            while len(en_par) and len(tr_par):
                # Create temporary documents
                temp_doc_1, temp_doc_2 = self.create_doc(2)

                count_1 = 0
                for paragraph in en_par:
                    temp_doc_1.add_paragraph(paragraph).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    if self.get_page_number(temp_doc_1) > 1:
                        break
                    count_1 += 1
                    
                count_2 = 0
                for paragraph in tr_par:
                    temp_doc_2.add_paragraph(paragraph).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    if self.get_page_number(temp_doc_2) > 1:
                        break
                    count_2 += 1
                
                par_count = max(min(count_1, count_2), 1)
                self.pars_per_page.append(par_count)
                
                del en_par[:par_count]
                del tr_par[:par_count]

                # Update progress
                progress.update(task, advance=1, description=f"[green]Calculating paragraph counts per page {progress_len-len(en_par)}/{progress_len}...")
                

    def create_doc(self, num=1):
        """  """
        docs = []
        for x in range(num):
            doc = Document()
            sec = doc.sections[0]
            sec.page_width = Inches(5.83)   # A5 width (14.8 cm)
            sec.page_height = Inches(8.27)  # A5 height (21 cm)
            sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1.5 / 2.54)
            
            docs.append(doc)

        return docs[0] if len(docs) == 1 else docs


    def get_page_number(self, doc):
        """  """
        doc.save("temp.docx")

        self.clear_console()
        convert("temp.docx", "temp.pdf", keep_active=True)
        self.clear_console()

        with open("temp.pdf", 'rb') as f:
            reader = PdfReader(f)
            page_num = len(reader.pages) 

        return page_num


    def remove_temps(self):
        """  """
        os.remove("temp.docx")
        os.remove("temp.pdf")


    def clear_console(self):
        os.system("cls" if os.name == "nt" else "clear")


